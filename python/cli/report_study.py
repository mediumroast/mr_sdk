#!/bin/env python3

import sys
import argparse
import configparser
import docx

from docx import Document
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.section import WD_ORIENT, WD_SECTION
from datetime import datetime
from mediumroast.api.high_level import Auth as authenticate
from mediumroast.api.high_level import Studies as study
from mediumroast.api.high_level import Interactions as interaction

### General utilities
def parse_cli_args(program_name='report_study', desc='A mediumroast.io utility that generates a Microsoft Word formatted report for a study.'):
    parser = argparse.ArgumentParser(prog=program_name, description=desc)
    parser.add_argument('--exclude_substudies', help="The names for the substudies to exclude in a comma separated list",
                        type=str, dest='exclude_substudies')
    parser.add_argument('--rest_url', help="The URL of the target REST server",
                        type=str, dest='rest_url', default='http://mr-01:3000')
    parser.add_argument('--guid', help="The GUID for the study to be reported on.",
                        type=str, dest='guid', required=True)
    parser.add_argument('--user', help="User name",
                        type=str, dest='user', default='foo')
    parser.add_argument('--secret', help="Secret or password",
                        type=str, dest='secret', default='bar')
    parser.add_argument('--config_file', help="The location to the configuration files",
                        type=str, dest='config_file', default='./reports.ini')
    cli_args = parser.parse_args()
    return cli_args


def read_config(conf_file='./reports.ini'):
    c = configparser.ConfigParser()
    c.read(conf_file)
    return c

def get_interaction_name(guid):
    """Get the interaction name by the GUID
    """
    interaction_ctl = interaction(credential)
    return interaction_ctl.get_name_by_guid(guid)[1]['interactionName']

def _create_header(doc_obj, conf, font_size=7):
    date_string = f'{datetime.now():%Y-%m-%d %H:%M}'
    s = doc_obj.sections[0]
    header = s.header
    header_p = header.paragraphs[0]
    header_p.text = conf['org'] + "\t | \t Created on: " + date_string
    style = doc_obj.styles['Header']
    font = style.font
    font.name = conf['font']
    font.size = Pt(font_size)
    header_p.style = doc_obj.styles['Header']


def _create_footer(doc_obj, conf, font_size=7):
    date_string = f'{datetime.now():%Y-%m-%d %H:%M}'
    s = doc_obj.sections[0]
    footer = s.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = conf['confidentiality'] + "\t | \t" + conf['copyright']
    style = doc_obj.styles['Footer']
    font = style.font
    font.name = conf['font']
    font.size = Pt(font_size)
    footer_p.style = doc_obj.styles['Footer']


def _create_cover_page(doc_obj, study, conf, logo_size=60, font_size=30):

    # Generics
    title_font_size = Pt(font_size)  # Title Font Size
    logo_size = Pt(font_size*2.5)

    # Organization name and logo
    logo = conf['logo']
    logo_title = doc_obj.add_paragraph().add_run()
    logo_title.add_picture(logo, height=logo_size)

    # Define the Cover Title Style
    org = conf['org']  # Organization
    title = "\n\nTitle: " + study['studyName']
    cover_title = doc_obj.add_paragraph(title)
    style = doc_obj.styles['Title']
    font = style.font
    font.name = conf['font']
    font.size = title_font_size
    cover_title.style = doc_obj.styles['Title']

    # Define the Subtitle content
    subtitle = "A " + org + " study report enabling attributable market insights."
    cover_subtitle = doc_obj.add_paragraph("")
    s = cover_subtitle.add_run(subtitle)
    subtitle_font = s.font
    subtitle_font.bold = True

    # Define the Author content
    author = "Mediumroast Barrista Robot"
    cover_author = doc_obj.add_paragraph("\nAuthor: ")
    a = cover_author.add_run(author)
    author_font = a.font
    author_font.bold = True

    # Define the Creation date content
    creation_date = f'{datetime.now():%Y-%m-%d %H:%M}'
    cover_date = doc_obj.add_paragraph("Creation Date: ")
    d = cover_date.add_run(creation_date)
    date_font = d.font
    date_font.bold = True

    # Add a page break
    doc_obj.add_page_break()


def _create_summary(doc_obj, study_doc, conf):
    # Create the Introduction section
    section_title = doc_obj.add_paragraph(
        'Findings')  # Create the Findings section
    section_title.style = doc_obj.styles['Title']
    doc_obj.add_heading('Introduction')
    clean_intro = " ".join(study_doc['Introduction'].split("\n"))
    doc_obj.add_paragraph(clean_intro)

    # Create the Opportunity section
    doc_obj.add_heading('Opportunity')
    clean_opportunity = " ".join(study_doc['Opportunity']['text'].split("\n"))
    doc_obj.add_paragraph(clean_opportunity)
    # Remove the text section before we process the numbered bullets
    del(study_doc['Opportunity']['text'])
    for opp in study_doc['Opportunity']:
        clean_opp = " ".join(study_doc['Opportunity'][opp].split("\n"))
        doc_obj.add_paragraph(clean_opp, style='List Bullet')

    # Create the Action section
    doc_obj.add_heading('Actions')
    clean_action = " ".join(study_doc['Action']['text'].split("\n"))
    doc_obj.add_paragraph(clean_action)
    # Remove the text section before we process the numbered bullets
    del(study_doc['Action']['text'])
    for action in study_doc['Action']:
        clean_act = " ".join(study_doc['Action'][action].split("\n"))
        doc_obj.add_paragraph(clean_act, style='List Number')

    # Add a page break
    doc_obj.add_page_break()


def _add_hyperlink(paragraph, text, url):
    """Taken from https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def _create_reference(interaction_guid, substudy, doc_obj, conf, char_limit=500):
    interaction_ctl = interaction(credential)
    success, interaction_data = interaction_ctl.get_by_guid(interaction_guid)
    if success:
        doc_obj.add_heading(interaction_data['interactionName'], 2)
        my_time = str(interaction_data['time'][0:2]) + \
            ':' + str(interaction_data['time'][2:4])
        my_date = str(interaction_data['date'][0:4]) + '-' + str(interaction_data['date'][4:6]) + '-' \
            + str(interaction_data['date'][6:8])
        interaction_meta = "\t\t|\t".join(['Date: ' + my_date + "\t" + my_time,
                                           'Sub-Study Identifier: ' + substudy])
        doc_obj.add_paragraph(interaction_meta)
        doc_obj.add_paragraph(
            interaction_data['abstract'][0:char_limit] + '...')
        resource = doc_obj.add_paragraph('Interaction Resource: ')
        _add_hyperlink(
            resource, interaction_data['interactionName'], interaction_data['url'].replace('s3', 'http'))
    else:
        print(
            'Something went wrong obtaining the interaction data for [' + interaction_guid + ']')


def _create_references(doc_obj, substudy_list, conf):
    section_title = doc_obj.add_paragraph(
        'References')  # Create the References section
    section_title.style = doc_obj.styles['Title']
    for substudy in substudy_list:
        for interaction in substudy_list[substudy]['interactions']:
            interaction_guid = substudy_list[substudy]['interactions'][interaction]['GUID']
            _create_reference(interaction_guid, substudy, doc_obj, conf)

def _create_quote(doc_obj, quote, indent, font_size):
    my_quote = quote
    my_para = doc_obj.add_paragraph(style='List Bullet')
    my_para.paragraph_format.left_indent = Pt(1.5 * indent)
    my_bullet = my_para.add_run(my_quote)
    my_bullet.font.size = Pt(font_size)
    my_para.paragraph_format.space_after = Pt(3)

def _create_quotes(doc_obj, quotes, indent, font_size, location='quotes'):
    for quote in quotes:
        my_quote = quotes[quote][location]
        my_para = doc_obj.add_paragraph(style='List Bullet')
        my_para.paragraph_format.left_indent = Pt(1.5 * indent)
        my_bullet = my_para.add_run(my_quote)
        my_bullet.font.size = Pt(font_size)
        my_para.paragraph_format.space_after = Pt(3)

def _create_subsection(doc_obj, start_text, body_text, indent, font_size, to_bold=False, to_italics=False):
    para = doc_obj.add_paragraph()
    para.paragraph_format.left_indent = Pt(indent)
    start_run = para.add_run(start_text)
    start_run.font.bold = to_bold
    start_run.font.size = Pt(font_size)
    body_run=para.add_run(body_text)
    body_run.font.size = Pt(font_size)
    if to_italics: body_run.font.italic = to_italics

def _create_intro(doc_obj, intro_name, intro_body, heading_level=2):
    doc_obj.add_heading(intro_name, level=heading_level)
    doc_obj.add_paragraph(intro_body)


def _create_key_theme(doc_obj, themes, quotes, conf, include_fortune=True):
    
    ### Define the summary theme
    _create_intro(doc_obj, 
        'Summary Theme', 
        conf['themes']['summary_intro'].replace("\n", " "))

    ## Create the definition
    theme = 'summary_theme'
    _create_subsection(doc_obj, 
            'Definition: ', 
            themes[theme]['description'], 
            int(conf['themes']['indent']), 
            font_size = int(conf['themes']['font_size']),
            to_bold = True)

    ## Determine if we should include the theme fortune or not
    if include_fortune:
        _create_subsection(doc_obj, 
            'Fortune: ', 
            themes[theme]['fortune'][0].upper() + themes[theme]['fortune'][1:] + ' [system generated]', 
            int(conf['themes']['indent']), 
            font_size = int(conf['themes']['font_size']),
            to_bold = True)
    
    ## Create the tags
    _create_subsection(doc_obj, 
            'Tags: ', 
            " | ".join(themes[theme]['tags'].keys()), 
            int(conf['themes']['indent']), 
            font_size = int(conf['themes']['font_size']),
            to_bold = True,
            to_italics = True)
    
    ## Create the quotes
    subsection_name = 'Theme Quotes'
    doc_obj.add_heading(subsection_name, level=3)
    _create_quotes(doc_obj, 
        quotes['summary'], 
        int(conf['themes']['indent']), 
        font_size = int(conf['themes']['font_size']))

    ### Add the discrete/detailed themes
    theme_loc = 'discrete_themes'
    quotes_loc = 'discrete'
    
    ## Create the starting paragraph
    _create_intro(doc_obj, 
        'Detailed Themes', 
        conf['themes']['discrete_intro'].replace("\n", " "))
    
    ## Add in the individual themes and their quotes
    my_themes = themes[theme_loc]
    for my_theme in my_themes:

        # Put in the theme identifier
        _create_intro(doc_obj, 
            'Detailed Theme Identifier: ' + my_theme, 
            conf['themes']['discrete_theme_intro'].replace("\n", " "),
            heading_level=3)
        
        # Add the description
        _create_subsection(doc_obj, 
            'Definition: ', 
            my_themes[my_theme]['description'], 
            int(conf['themes']['indent']), 
            font_size = int(conf['themes']['font_size']),
            to_bold = True)

        # Include the fortune if the setting is true
        if include_fortune:
            _create_subsection(doc_obj, 
                'Fortune: ', 
                my_themes[my_theme]['fortune'][0].upper() + my_themes[my_theme]['fortune'][1:] + ' [system generated]', 
                int(conf['themes']['indent']), 
                font_size = int(conf['themes']['font_size']),
                to_bold = True)

        # Add the tags
        _create_subsection(doc_obj, 
            'Tags: ', 
            " | ".join(my_themes[my_theme]['tags'].keys()), 
            int(conf['themes']['indent']), 
            font_size = int(conf['themes']['font_size']),
            to_bold = True,
            to_italics = True)

        # Pull in the quotes
        subsection_name = 'Theme Quotes by Interaction'
        doc_obj.add_heading(subsection_name, level=4)
        if my_theme in quotes[quotes_loc]:
            for interaction in quotes[quotes_loc][my_theme]:
                doc_obj.add_heading(get_interaction_name(interaction), level=5)
                the_quotes = quotes[quotes_loc][my_theme][interaction]['quotes']
                # Explain that the system was not able to find a relevant quote
                if not the_quotes: the_quotes=[['mediumroast.io was unable to find a relevant quote or text snippet for this theme.']]
                for my_quote in the_quotes:
                    _create_quote(doc_obj, 
                        my_quote[0], 
                        int(conf['themes']['indent']), 
                        font_size = int(conf['themes']['font_size']))
                
                _create_subsection(doc_obj, 
                    'Frequency: ', 
                    str(quotes[quotes_loc][my_theme][interaction]['frequency']), 
                    int(conf['themes']['indent']), 
                    font_size = int(conf['themes']['font_size']),
                    to_bold = True,
                    to_italics = True)
    
    doc_obj.add_page_break()


def _create_key_themes(doc_obj, substudies, substudy_excludes, conf):
    section_title = doc_obj.add_paragraph(
        'Key Themes by Sub-Study')  # Create the Themes section
    section_title.style = doc_obj.styles['Title']
    doc_obj.add_paragraph(conf['themes']['intro'].replace("\n", " "))
    for substudy in substudies:
        if substudy in substudy_excludes:
            continue
        doc_obj.add_heading('Sub-Study Identifier: ' + substudy + ' — ' + substudies[substudy]['description'], 1)
        _create_key_theme(
            doc_obj, substudies[substudy]['keyThemes'], substudies[substudy]['keyThemeQuotes'], conf)

def change_orientation(doc_obj):
    current_section = doc_obj.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = doc_obj.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

    return new_section

def _create_row(the_row, id, type,freq, snip, src):
    ID = 0
    TYPE = 1 
    FREQ = 2
    SNIP = 3
    SRC = 5
    the_row[ID].text = str(id)
    the_row[TYPE].text = str(type)
    the_row[FREQ].text = str(freq)
    the_row[SNIP].text = str(snip)
    the_row[SRC].text = str(src)


def _create_rows():
    """
    For summary
    create single row

    For discrete
        foreach theme
            create single row 
    """
    pass

def _create_summary_theme_tables(doc_obj, substudies, substudy_excludes, conf):
    change_orientation(doc_obj) # Flip to landscape mode
    section_title = doc_obj.add_paragraph(
        'Key Theme Summary Tables')  # Create the References section
    section_title.style = doc_obj.styles['Title']
    for substudy in substudies:
        if substudy in substudy_excludes:
            continue
        doc_obj.add_heading('Sub-Study Identifier: ' + substudy + ' — ' + substudies[substudy]['description'], 1)
        my_table = doc_obj.add_table(rows=1, cols=6)
        my_table.style = 'Dark List'
        header_row = my_table.rows[0].cells
        header_row[0].text = 'Identifier'
        header_row[1].text = 'Type'
        header_row[2].text = 'Frequency'
        header_row[3].text = 'Snippet'
        header_row[4].text = 'Source'
        my_row = my_table.add_row().cells

        ## Process the summary theme
        my_theme = 'Summary Theme'
        my_type = 'Summary'
        my_frequency = 'N/A'
        my_interaction = list(substudies[substudy]['keyThemeQuotes']['summary'].keys())[0]
        my_snippet = substudies[substudy]['keyThemeQuotes']['summary'][my_interaction]['quotes'][0]
        my_source = get_interaction_name(my_interaction)
        _create_row(my_row, my_theme, my_type, my_frequency, my_snippet, my_source)

        doc_obj.add_page_break()

    change_orientation(doc_obj) # Flip to portrait mode





def report(study, conf, substudy_excludes):
    # Document generics
    d = Document()  # Create doc object
    style = d.styles['Normal']
    font = style.font
    font.name = conf['font']
    font.size = Pt(int(conf['font_size']))
    _create_cover_page(d, study, conf)  # Create the cover page
    _create_header(d, conf)  # Create the doc header
    _create_footer(d, conf)  # Create the doc footer
    
    ### Intro, opportunity and actions sections
    _create_summary(d, study['document'], conf)

    ### Key Themes
    ## Key Themes Summary Table
    _create_summary_theme_tables(d, study['substudies'], substudy_excludes, conf)
    
    ## Detailed Key Themes
    _create_key_themes(d, study['substudies'], substudy_excludes, conf)
    
    ### References
    _create_references(d, study['substudies'], conf)

    return d


if __name__ == "__main__":
    my_args = parse_cli_args()
    configurator = read_config(conf_file=my_args.config_file)

    # Set default items from the configuration file for the report
    report_conf = {
        'org': configurator['DEFAULT']['organization_name'],
        'logo': configurator['DEFAULT']['logo_image'],
        'font': configurator['DEFAULT']['font_type'],
        'font_size': configurator['DEFAULT']['font_size'],
        'font_measure': configurator['DEFAULT']['font_measure'],
        'copyright': configurator['DEFAULT']['copyright_notice'],
        'confidentiality': configurator['DEFAULT']['confidential_notice'],
        'themes': {
            'font_size': configurator['THEME_FORMAT']['font_size'],
            'intro': configurator['THEME_FORMAT']['key_theme_intro'],
            'summary_intro': configurator['THEME_FORMAT']['summary_theme_intro'],
            'discrete_intro': configurator['THEME_FORMAT']['discrete_themes_intro'],
            'discrete_theme_intro': configurator['THEME_FORMAT']['discrete_theme_intro'],
            'indent': configurator['THEME_FORMAT']['indent'],
        }
    }

    auth_ctl = authenticate(
        user_name=my_args.user, secret=my_args.secret, rest_server_url=my_args.rest_url)
    credential = auth_ctl.login()
    substudy_excludes = my_args.exclude_substudies.split(',')
    study_ctl = study(credential)
    success, study_obj = study_ctl.get_by_guid(my_args.guid)
    if success:
        doc_name = study_obj['studyName'].replace(
            ' ', '_') + "_study_report.docx"
        document = report(study_obj, report_conf, substudy_excludes)
        document.save(doc_name)

    else:
        print('CLI ERROR: This is a generic error message, as something went wrong.')
        sys.exit(-1)
